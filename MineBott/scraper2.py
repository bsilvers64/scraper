from bs4 import BeautifulSoup
import requests
import docx

def getpage(url):        #returns a bs4 object of any url passed
    # added header for fake user-agent
    # replace this with user-agent of your system. find it at - https://www.whatismybrowser.com/detect/what-is-my-user-agent
    headers = {
        "User-Agent": 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.130 Safari/537.36 OPR/66.0.3515.44'}
    try:
        res = requests.get(url, headers=headers)
    except requests.ConnectionError as e:
        print('Connection error occoured')
        return None
    except requests.RequestException as e:
        print('error occoured')
        return None
    return BeautifulSoup(res.text, 'lxml')


def option_a():         # nseindia
    doc = docx.Document()
    f = open("corporate_data.text", '+w')
    doc.add_paragraph("the data is in this format : -\n(Symbol)---(Company Name)---(Series)---(Face Value)---(Purpose)---(Ex-Date)---(Record Date)---(BC Start-Date)---(BC End-Date)---(ND Start Date)---(ND End Date)")
    for pagecount in range(2,25):
        url = 'https://www1.nseindia.com/sme/marketinfo/corporates/actions/latestCorpActions.jsp?currentPage='         # url
        url += str(pagecount)
        page = getpage(url)      #bs4 object
        row = page.find_all('tr', class_='alt')
        for i in row:
            td = i.find_all('td', {'class': ["normaltext", "date"]})
            comma = doc.add_paragraph()
            comma.add_run("\n")
            f.write("\n")
            for j in range((len(list(td)))):
                com = td[j].get_text()
                comma.add_run(com+" , ")
                f.write(com+" , ")
    f.close()
    doc.save("corp_data.docx")
    print("scrapping successful !!")


def option_b():         # moneycontrol
    url = 'https://www.moneycontrol.com/stocks/marketinfo/dividends_declared/homebody.php?sel_year='
    # change the year to anything from 2000 to 2020
    year = int(input("enter the year for data query [2000 to 2020] : \n"))
    url += str(year)
    page = getpage(url) #bs4 object
    data_table = page.find('table', class_='b_12 dvdtbl')
    data_rows = data_table.find_all('tr')
    data_rows = data_rows[2:len(data_rows)]
    f = open("corporate_data.text", '+w')
    g = open("corporate_data.csv", '+w')
    doc = docx.Document()
    comma = doc.add_paragraph('the data is in this format : -\n(Company-name)---(Dividend type)---(Dividend %)---(Date announced)---(Date recorded)---(Date Ex-dividend)\n\n')
    f.write('the data is in this format : -\n(Company name)---(Dividend type)---(Dividend %)---(Date announced)---(Date recorded)---(Date Ex-dividend)\n\n')

    for i in data_rows:
        row = i.find_all('td')
        comma.add_run('\n\n')
        f.write('\n\n')
        for j in row:
            f.write(j.get_text("|", strip='true')+"  ,  ")
            g.write(j.get_text("|", strip='true') + "  ,  ")
            comma.add_run(j.get_text("|", strip='true')+"  ,  ")
        g.write('\n')

    g.close()
    f.close()
    doc.save('corp_data.docx')
    print("scrapping successful !!")

def main():
    s = str(input("press 'a' for source 1 (nseindia.com)  \nand 'b' for source 2 (moneycontrol.com) : "))
    if s == 'a':
        option_a()
    elif s == 'b':
        option_b()
    else:
        print("type the correct option number")

if __name__== "__main__":
    main()