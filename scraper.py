from bs4 import BeautifulSoup
import requests
import docx

doc = docx.Document()

# urls
url = 'https://www1.nseindia.com/sme/marketinfo/corporates/actions/latestCorpActions.jsp?currentPage'

#added header for fake user-agent
#replace this with user-agent of your system. find it at - https://www.whatismybrowser.com/detect/what-is-my-user-agent
headers = { #enter user-agent here }

f = open("compdata.text", '+a')
title = doc.add_paragraph("the data is in this format : -\n(Symbol)---(Company Name)---(Series)---(Face Value)---(Purpose)---(Ex-Date)---(Record Date)---(BC Start-Date)---(BC End-Date)---(ND Start Date)---(ND End Date)")

for pagecount in range(2,11):
    url+="="+str(pagecount)
    try :
        res = requests.get(url,headers=headers)
    except requests.ConnectionError as e:
        print(e)

    #bs4 object
    page = BeautifulSoup(res.text,'lxml')
    row = page.find_all('tr',class_='alt')

    for i in row :
        td = i.find_all('td', {'class': ["normaltext", "date"]})
        comma = doc.add_paragraph()
        comma.add_run("\n")
        f.write("\n")
        for j in range((len(list(td)))):
            com = (td)[j].get_text()
            comma.add_run(com+" , ")
            f.write(com+" , ")


f.close()
doc.save("comp_data.docx")
print("success !! data scraped")
